import re
import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_paragraph(doc, text="", style='Normal', space_after=6, space_before=0, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    return p

def format_run(run, font_name="Times New Roman", font_size=12, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def parse_inline_markdown(paragraph, text, font_name="Times New Roman", font_size=12, base_color=None):
    # Regex to handle **bold**, *italic*, `code`, math $...$
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\$.*?\$|\[.*?\]\(.*?\))')
    tokens = pattern.split(text)
    
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            content = token[2:-2]
            r = paragraph.add_run(content)
            format_run(r, font_name=font_name, font_size=font_size, bold=True, color=base_color)
        elif token.startswith('*') and token.endswith('*'):
            content = token[1:-1]
            r = paragraph.add_run(content)
            format_run(r, font_name=font_name, font_size=font_size, italic=True, color=base_color)
        elif token.startswith('`') and token.endswith('`'):
            content = token[1:-1]
            r = paragraph.add_run(content)
            format_run(r, font_name="Consolas", font_size=11.5, color=RGBColor(199, 37, 78))
        elif token.startswith('$') and token.endswith('$'):
            content = token.strip('$')
            r = paragraph.add_run(content)
            format_run(r, font_name="Times New Roman", font_size=font_size, italic=True, color=base_color)
        elif token.startswith('[') and ']' in token:
            match = re.match(r'\[(.*?)\]\((.*?)\)', token)
            if match:
                link_text, link_url = match.groups()
                r = paragraph.add_run(link_text)
                format_run(r, font_name=font_name, font_size=font_size, color=RGBColor(0, 102, 204), italic=True)
            else:
                r = paragraph.add_run(token)
                format_run(r, font_name=font_name, font_size=font_size, color=base_color)
        else:
            r = paragraph.add_run(token)
            format_run(r, font_name=font_name, font_size=font_size, color=base_color)

def build_docx(md_filepath, docx_filepath):
    doc = Document()
    
    # Configure Page Margins: Left 4cm (1.57 inches), Top/Bottom/Right 1 inch (2.54 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.57) # 4cm for binding as requested in UMaT guidelines
        section.right_margin = Inches(1.0)

    # Set base Normal style font
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)

    with open(md_filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    in_title_page = True
    in_toc = False

    TOC_PAGE_MAP = {
        "Abstract": "ii",
        "Acknowledgements": "iii",
        "Dedication": "iv",
        "Keywords": "v",
        "Chapter 1: Introduction, Aims and Objectives": "1",
        "1.1 Introduction to Problem": "1",
        "1.2 Introduction to Project, Aim and Objectives": "1",
        "1.2.1 Project Overview": "1",
        "1.2.2 Aim": "2",
        "1.2.3 Objectives": "2",
        "1.3 Research Questions": "3",
        "1.4 Scope of the Project": "3",
        "1.5 Project Justification": "4",
        "1.6 Organization of Chapters": "4",
        "Chapter 2: Literature Review": "5",
        "2.1 Theoretical Foundations of API Security": "5",
        "2.2 Review of Security Taxonomies: OWASP, MITRE ATT&CK, and CWE": "6",
        "2.2.1 OWASP API Security Top 10:2023 Taxonomy": "6",
        "2.2.2 MITRE ATT&CK Knowledge Base Alignment": "7",
        "2.2.3 Common Weakness Enumeration (CWE) Mapping": "7",
        "2.3 Evaluation of Existing API Security Tools and Operational Gaps": "8",
        "2.3.1 Analysis of APIsec Surface Suite": "8",
        "2.3.2 Analysis of Tooling Gaps and MazAPI Differentiation": "9",
        "2.5 Static Secret Analysis and Browser Interception Techniques": "10",
        "2.5.1 Static Code Secret Analysis": "10",
        "2.5.2 Dynamic Browser Interception via Manifest V3": "11",
        "Chapter 3: Methodology": "12",
        "3.1 Agile-Iterative Engineering Framework": "12",
        "3.2 System Architecture and Enterprise Production Deployment Stack": "13",
        "3.3 Transparent Monitoring Proxy and Rule-Based Pre-Check Design": "14",
        "3.3.1 Rule-Based BOLA Pre-Check Layer": "14",
        "3.3.2 Dynamic OpenAPI 3.0 Synthesizer, Schema Drift & Active Inline Auto-Blocking": "15",
        "3.3.3 Unified BOM Generator & Model Context Protocol (MCP) Auditor": "16",
        "3.4 Feature Engineering and Dataset Synthesizer Specification": "17",
        "3.4.1 Dataset Synthesizer Implementation": "17",
        "3.5 Machine Learning Ensemble Architecture": "18",
        "3.6 MazAPI Web Scanner and Playwright Session Interception Engine": "19",
        "3.6.1 Endpoint Discovery Strategies": "19",
        "3.6.2 Playwright Headless Session Interception": "20",
        "3.7 Manifest V3 Browser Extension Architecture": "21",
        "3.7.1 Service Worker Security Probe Execution": "21",
        "3.8 Visual Studio Code Static Analysis Extension Engineering": "22",
        "3.8.1 Multi-Layer Detection Strategy": "22",
        "3.9 Interactive Command-Line Management Console": "23",
        "Chapter 4: Design, Testing and Evaluation": "24",
        "4.1 Comparative Vulnerable vs. Hardened API Implementation": "24",
        "4.1.1 Implementation Comparison Matrix": "25",
        "4.2 Empirical Security Testing Results across Vulnerability Classes": "26",
        "4.2.1 Detailed Evaluation Breakdown": "26",
        "4.3 Machine Learning Ensemble Performance Metrics and Evaluation": "28",
        "4.3.1 Classification Performance Metrics": "28",
        "4.3.2 Confusion Matrix Analysis": "29",
        "4.3.3 Feature Importance Ranking": "30",
        "4.3.4 Operational Latency Overhead Analysis": "31",
        "4.4 Validation on External Real-World Targets and VulnBank Lab": "32",
        "4.4.1 Google Gemini API External Scanning Validation": "32",
        "4.4.2 VulnBank Banking Lab Evaluation": "33",
        "4.4.3 API Surface OpenAPI Documentation Comparison": "34",
        "4.5 VS Code Extension Secret Scanning and Static Analysis Benchmark": "35",
        "4.6 External Attack Workflow Validation using Kali Linux": "36",
        "Chapter 5: Conclusions & Further Work": "38",
        "5.1 Summary of Findings and Contributions": "38",
        "5.2 System Limitations and Challenges": "39",
        "5.3 Recommendations for Future Work": "40",
        "References": "41",
        "Appendices": "42",
        "Appendix A: Implementation Schedule and Project Gantt Chart": "42",
        "Appendix B: Vulnerability and Defense Mapping Matrix": "43",
        "Appendix C: Feature Vector Pipeline and ML Model Parameters": "44",
        "Appendix D: Static Secret Detector Regex Patterns and Compliance Mappings": "45"
    }

    def flush_table(lines):
        if not lines:
            return
        # Parse table rows
        rows_data = []
        for l in lines:
            if re.match(r'^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$', l):
                continue # Header separator line
            parts = [p.strip() for p in l.strip().strip('|').split('|')]
            rows_data.append(parts)
        if not rows_data:
            return
        
        num_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for i, row in enumerate(rows_data):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    cell = table.cell(i, j)
                    cell.text = ""
                    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.line_spacing = 1.15
                    
                    if i == 0: # Header Row
                        set_cell_background(cell, "1F497D") # Dark Blue
                        parse_inline_markdown(p, cell_text, font_name="Times New Roman", font_size=12, base_color=RGBColor(255, 255, 255))
                    else: # Data Rows
                        bg = "F2F5F9" if i % 2 == 1 else "FFFFFF"
                        set_cell_background(cell, bg)
                        parse_inline_markdown(p, cell_text, font_name="Times New Roman", font_size=12)
        
        # Add spacing after table
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def flush_code_block(lines):
        if not lines:
            return
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "F8F9FA")
        set_cell_margins(cell, top=150, bottom=150, left=200, right=200)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        
        code_text = "".join(lines)
        r = p.add_run(code_text)
        format_run(r, font_name="Consolas", font_size=11.5, color=RGBColor(40, 40, 40))
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Handle Code Blocks
        if line.startswith('```'):
            if in_code_block:
                flush_code_block(code_block_lines)
                code_block_lines = []
                in_code_block = False
            else:
                if in_table:
                    flush_table(table_lines)
                    table_lines = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Handle Tables
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table(table_lines)
                table_lines = []
                in_table = False

        # Blank Lines
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # Horizontal Rule
        if stripped in ['---', '***', '___']:
            in_toc = False
            p = add_styled_paragraph(doc, space_after=18, space_before=12)
            r = p.add_run("_________________________________________________________________________________")
            format_run(r, font_size=9, color=RGBColor(180, 180, 180))
            i += 1
            continue

        # Headings
        if stripped.startswith('# '):
            p = add_styled_paragraph(doc, space_after=18, space_before=24, align=WD_ALIGN_PARAGRAPH.CENTER)
            parse_inline_markdown(p, stripped[2:], font_size=18, base_color=RGBColor(0, 0, 0))
            for run in p.runs:
                run.bold = True
        elif stripped.startswith('## '):
            heading_text = stripped[3:]
            if heading_text == "Abstract":
                in_title_page = False
            
            if heading_text == "Table of Contents":
                in_toc = True
            else:
                in_toc = False

            if in_title_page:
                p = add_styled_paragraph(doc, space_after=18, space_before=24)
                parse_inline_markdown(p, heading_text, font_size=15, base_color=RGBColor(0, 0, 0))
                for run in p.runs:
                    run.bold = True
            else:
                if heading_text.startswith('Chapter') or heading_text.startswith('References') or heading_text.startswith('Appendices'):
                    p = doc.add_paragraph()
                    p.add_run().add_break(WD_BREAK.PAGE)
                p = add_styled_paragraph(doc, style='Heading 1', space_after=18, space_before=24)
                parse_inline_markdown(p, heading_text, font_size=15, base_color=RGBColor(0, 0, 0))
                for run in p.runs:
                    run.bold = True
        elif stripped.startswith('### '):
            heading_text = stripped[4:]
            if in_title_page:
                p = add_styled_paragraph(doc, space_after=12, space_before=18)
                parse_inline_markdown(p, heading_text, font_size=13, base_color=RGBColor(0, 0, 0))
                for run in p.runs:
                    run.bold = True
            else:
                p = add_styled_paragraph(doc, style='Heading 2', space_after=12, space_before=18)
                parse_inline_markdown(p, heading_text, font_size=13, base_color=RGBColor(0, 0, 0))
                for run in p.runs:
                    run.bold = True
        elif stripped.startswith('#### '):
            heading_text = stripped[5:]
            if in_title_page:
                p = add_styled_paragraph(doc, space_after=10, space_before=12)
                parse_inline_markdown(p, heading_text, font_size=12, base_color=RGBColor(0, 0, 0))
                for run in p.runs:
                    run.bold = True
                    run.italic = True
            else:
                p = add_styled_paragraph(doc, style='Heading 3', space_after=10, space_before=12)
                parse_inline_markdown(p, heading_text, font_size=12, base_color=RGBColor(0, 0, 0))
                for run in p.runs:
                    run.bold = True
                    run.italic = True
        # Bullet list items
        elif stripped.startswith('- ') or stripped.startswith('* ') or re.match(r'^\d+\.\s', stripped):
            if in_toc:
                leading_spaces = len(line) - len(line.lstrip())
                if leading_spaces >= 4:
                    left_indent = Inches(0.50)
                    font_size = 11
                    is_bold = False
                elif leading_spaces >= 2:
                    left_indent = Inches(0.25)
                    font_size = 11.5
                    is_bold = False
                else:
                    left_indent = Inches(0.0)
                    font_size = 12
                    is_bold = True

                match_link = re.search(r'\[(.*?)\]\(.*?\)', stripped)
                if match_link:
                    title_text = match_link.group(1).strip()
                else:
                    title_text = stripped[2:].strip()

                page_num = TOC_PAGE_MAP.get(title_text, "1")

                p = add_styled_paragraph(doc, space_after=4, space_before=4, line_spacing=1.15)
                p.paragraph_format.left_indent = left_indent
                p.paragraph_format.tab_stops.add_tab_stop(Inches(5.93), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)

                r_text = p.add_run(f"{title_text}\t")
                format_run(r_text, font_size=font_size, bold=is_bold, color=RGBColor(0, 0, 0))

                r_page = p.add_run(page_num)
                format_run(r_page, font_size=font_size, bold=is_bold, color=RGBColor(0, 0, 0))
            else:
                p = add_styled_paragraph(doc, space_after=8, line_spacing=1.5)
                p.paragraph_format.left_indent = Inches(0.25)
                
                if stripped.startswith('- ') or stripped.startswith('* '):
                    r_bullet = p.add_run("• ")
                    format_run(r_bullet, font_size=12, bold=True)
                    content = stripped[2:]
                else:
                    match = re.match(r'^(\d+\.)\s*(.*)', stripped)
                    r_bullet = p.add_run(match.group(1) + " ")
                    format_run(r_bullet, font_size=12, bold=True)
                    content = match.group(2)
                
                parse_inline_markdown(p, content, font_size=12)
        # Blockquote / Figure Caption / Standard Paragraph
        elif stripped.startswith('> '):
            p = add_styled_paragraph(doc, space_after=10, space_before=6, line_spacing=1.3)
            p.paragraph_format.left_indent = Inches(0.4)
            parse_inline_markdown(p, stripped[2:], font_size=11, base_color=RGBColor(80, 80, 80))
            for run in p.runs:
                run.italic = True
        elif stripped.startswith('!['): # Image placeholder / caption
            match = re.match(r'^!\[(.*?)\]\((.*?)\)', stripped)
            if match:
                caption, img_rel_path = match.groups()
                md_dir = os.path.dirname(os.path.abspath(md_filepath))
                img_path = os.path.join(md_dir, img_rel_path.replace('/', os.sep))
                
                if os.path.exists(img_path):
                    # Add centered image
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(12)
                    p_img.paragraph_format.space_after = Pt(6)
                    r_img = p_img.add_run()
                    r_img.add_picture(img_path, width=Inches(5.5))
                    
                    # Add centered caption
                    p_cap = add_styled_paragraph(doc, space_after=12, space_before=4, align=WD_ALIGN_PARAGRAPH.CENTER)
                    r_cap = p_cap.add_run(f"Figure: {caption}")
                    format_run(r_cap, font_size=10.5, italic=True, color=RGBColor(100, 100, 100))
                else:
                    p = add_styled_paragraph(doc, space_after=12, space_before=6, align=WD_ALIGN_PARAGRAPH.CENTER)
                    parse_inline_markdown(p, stripped, font_size=10.5, base_color=RGBColor(100, 100, 100))
                    for run in p.runs:
                        run.italic = True
            else:
                p = add_styled_paragraph(doc, space_after=12, space_before=6, align=WD_ALIGN_PARAGRAPH.CENTER)
                parse_inline_markdown(p, stripped, font_size=10.5, base_color=RGBColor(100, 100, 100))
                for run in p.runs:
                    run.italic = True
        else:
            p = add_styled_paragraph(doc, space_after=10, space_before=0, line_spacing=1.5)
            parse_inline_markdown(p, stripped, font_size=12)

        i += 1

    try:
        doc.save(docx_filepath)
        print(f"Document successfully created: {docx_filepath}")
    except PermissionError:
        backup_path = docx_filepath.replace(".docx", "_output.docx")
        doc.save(backup_path)
        print(f"WARNING: Permission denied when writing to {docx_filepath} (likely open in Word).")
        print(f"Saved document to backup path: {backup_path}")

if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, "MazAPI_Final_Report.md")
    docx_path = os.path.join(script_dir, "MazAPI_Final_Report.docx")
    build_docx(md_path, docx_path)
